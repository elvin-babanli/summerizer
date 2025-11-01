
from __future__ import annotations
from flask_session import Session
import os
import io
import re
from datetime import datetime
from typing import List, Any
from flask import (
    Flask, render_template, request, redirect, url_for,
    send_file, session, flash, jsonify, current_app
)
from werkzeug.utils import secure_filename

# CSRF
from flask_wtf import CSRFProtect

# Rate limit (optional but supported)
try:
    from flask_limiter import Limiter
    from flask_limiter.util import get_remote_address
except Exception:
    Limiter = None
    def get_remote_address():  # type: ignore
        return "127.0.0.1"

# Processing service
from services.summarizer import (
    GenerateOptions,
    FileAnalyzer,
    SummarizerService,
    build_base_filename,
)

# Optional exports
try:
    from docx import Document  # python-docx
except Exception:
    Document = None  # type: ignore

# PDF export (Unicode)
try:
    from reportlab.pdfgen import canvas  # type: ignore
    from reportlab.pdfbase import pdfmetrics  # type: ignore
    from reportlab.pdfbase.ttfonts import TTFont  # type: ignore
except Exception:
    canvas = None  # type: ignore
    pdfmetrics = None  # type: ignore
    TTFont = None  # type: ignore


# -----------------------------------------------------------------------------
# Flask App Factory
# -----------------------------------------------------------------------------

def create_app() -> Flask:
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
    app = Flask(
        __name__,
        template_folder=os.path.join(BASE_DIR, "templates"),
        static_folder=os.path.join(BASE_DIR, "static"),
    )

    # --------------------
    # Configuration
    # --------------------
    app.config["SECRET_KEY"] = os.getenv("SECRET_KEY", "change-me")
    app.config["UPLOAD_FOLDER"] = os.getenv("UPLOAD_FOLDER", os.path.join(os.getcwd(), "uploads"))
    app.config["OUTPUT_FOLDER"] = os.getenv("OUTPUT_FOLDER", os.path.join(os.getcwd(), "outputs"))
    app.config["MAX_CONTENT_LENGTH"] = 50 * 1024 * 1024  # 50MB per request
    app.config["MAX_TOTAL_PAGES"] = int(os.getenv("MAX_TOTAL_PAGES", "250"))
    app.config["TEMPLATES_AUTO_RELOAD"] = True
    app.config["SESSION_TYPE"] = "filesystem"
    app.config["SESSION_FILE_DIR"] = os.path.join(BASE_DIR, ".flask_session")
    app.config["SESSION_PERMANENT"] = False
    app.config["SESSION_USE_SIGNER"] = True
    app.config["SESSION_COOKIE_SAMESITE"] = "Lax"
    app.config["SESSION_COOKIE_SECURE"] = False  # dev mühitdə False, prod-da True

    # Session(app)

    # Ensure folders exist
    os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)
    os.makedirs(app.config["OUTPUT_FOLDER"], exist_ok=True)
    os.makedirs(app.config["SESSION_FILE_DIR"], exist_ok=True)


    # CSRF
    csrf = CSRFProtect()
    csrf.init_app(app)

    # Rate limiter (keep behavior if available)
    if Limiter is not None:
        limiter = Limiter(get_remote_address, app=app, default_limits=["120 per hour"])
    else:
        limiter = None  # type: ignore

    # --------------------
    # Helpers
    # --------------------
    ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}

    def _allowed_file(filename: str) -> bool:
        _, ext = os.path.splitext(filename.lower())
        return ext in ALLOWED_EXTENSIONS

    def _register_pdf_font_if_available(c: Any) -> None:
        """Register DejaVuSans if present for Unicode PDF output."""
        if pdfmetrics is None or TTFont is None:
            return
        candidate_paths: List[str] = [
            os.path.join("static", "fonts", "DejaVuSans.ttf"),
            os.path.join(os.getcwd(), "static", "fonts", "DejaVuSans.ttf"),
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        ]
        for p in candidate_paths:
            if os.path.isfile(p):
                try:
                    pdfmetrics.registerFont(TTFont("DejaVuSans", p))
                    c.setFont("DejaVuSans", 11)
                    return
                except Exception:
                    continue


    def _pdf_write_multiline(c: Any, text: str, x: int = 50, y_start: int = 800, line_height: int = 16, right_margin: int = 50) -> None:
        """
        Wrap lines by visible width (points) so text fits the page.
        """
        # Page metrics
        page_width, _ = c._pagesize  # e.g., 595 for A4 width (points)
        max_width = page_width - x - right_margin

        # Current font info (fallbacks)
        font_name = getattr(c, "_fontname", "DejaVuSans")
        font_size = getattr(c, "_fontsize", 11)

        def wrap_by_width(line: str):
            if not line:
                yield ""
                return
            parts = re.split(r"(\s+)", line)  # keep spaces
            buf = ""
            buf_w = 0.0
            for token in parts:
                w = pdfmetrics.stringWidth(token, font_name, font_size)
                if buf and (buf_w + w) > max_width:
                    yield buf.rstrip()
                    buf = token
                    buf_w = w
                else:
                    buf += token
                    buf_w += w
            if buf:
                yield buf.rstrip()

        y = y_start
        for raw_line in text.splitlines():
            for phys_line in wrap_by_width(raw_line):
                if y < 40:  # new page
                    c.showPage()
                    # re-apply font on new page
                    try:
                        c.setFont(font_name, font_size)
                    except Exception:
                        pass
                    y = y_start
                c.drawString(x, y, phys_line)
                y -= line_height

    # -----------------------------------------------------------------------------
    # Routes
    # -----------------------------------------------------------------------------

    @app.get("/")
    def index():
        # limits for UI
        limits = {
            "max_files": 10,
            "max_file_mb": 50,
            "max_total_pages": 250,
        }

        # list uploaded files (basic metadata)
        uploads_dir = current_app.config["UPLOAD_FOLDER"]
        files = []
        if os.path.isdir(uploads_dir):
            for name in sorted(os.listdir(uploads_dir)):
                path = os.path.join(uploads_dir, name)
                if not os.path.isfile(path):
                    continue
                ext = os.path.splitext(name)[1].lower()
                size_bytes = os.path.getsize(path)
                files.append({
                    "id": name,  # used as fid in template
                    "name": name,
                    "ext": ext[1:].upper() if ext else "",
                    "pages": 0,  # optional: real page count can be added later
                    "size_bytes": size_bytes,
                })

        # options defaults
        options = {
            "task": "summary",
            "words": 800,
            "language": "English",
            "output": "txt",
            "notes": "",
        }
        if "options" in session:
            options.update(session["options"])

        languages = [
            "English","Polish","Turkish","Azerbaijani","Russian","German",
            "French","Spanish","Italian","Portuguese","Ukrainian","Arabic",
            "Chinese","Japanese","Korean","Hindi"
        ]

        stats = {
            "files": len(files),
            "pages": 0,
            "bytes": sum(f["size_bytes"] for f in files)
        }

        result_text = session.get("last_result_text", "")

        return render_template(
            "index.html",
            limits=limits,
            files=files,
            options=options,
            languages=languages,
            stats=stats,
            result_text=result_text
        )

    @app.post("/upload")
    def upload():
        if "files" not in request.files:
            flash("No files part", "error")
            return redirect(url_for("index"))

        files = request.files.getlist("files")
        saved = 0
        for f in files:
            if not f or f.filename == "":
                continue
            if not _allowed_file(f.filename):
                continue
            fname = secure_filename(f.filename)
            f.save(os.path.join(current_app.config["UPLOAD_FOLDER"], fname))
            saved += 1

        if saved == 0:
            flash("No valid files uploaded.", "warning")
        else:
            flash(f"{saved} file(s) uploaded.", "success")
        return redirect(url_for("index"))

    @app.post("/remove/<path:fid>")
    def remove(fid):
        uploads_dir = current_app.config["UPLOAD_FOLDER"]
        path = os.path.join(uploads_dir, fid)
        try:
            if os.path.isfile(path):
                os.remove(path)
                flash("File removed.", "success")
            else:
                flash("File not found.", "warning")
        except Exception:
            flash("Could not remove file.", "error")
        return redirect(url_for("index"))

    @app.post("/reset")
    def reset():
        # Clear session
        session.clear()
        # Clean uploads dir
        uploads_dir = current_app.config["UPLOAD_FOLDER"]
        try:
            if os.path.isdir(uploads_dir):
                for name in os.listdir(uploads_dir):
                    path = os.path.join(uploads_dir, name)
                    if os.path.isfile(path):
                        try:
                            os.remove(path)
                        except Exception:
                            pass
        except Exception:
            pass
        flash("Session and uploads reset.", "success")
        return redirect(url_for("index"))

    # NOTE: Generate NOW updates the UI result first. Download happens only via /export.
    if limiter:
        @app.post("/generate")
        @limiter.limit("6 per minute")
        def generate():
            return _generate_impl()
    else:
        @app.post("/generate")
        def generate():
            return _generate_impl()

    def _generate_impl():
        """
        - Read task/words/language/output/notes
        - Build corpus from uploaded files
        - Build prompt and call API (or mock)
        - Save ONLY to session and redirect to index (show in Result editor)
        - Download happens later via /export
        """
        task = request.form.get("task", "summary")
        words = int((request.form.get("words") or "800").strip() or 800)
        language = request.form.get("language", "English").strip()
        output = request.form.get("output", "txt").strip().lower()  # txt | docx | pdf
        notes = request.form.get("notes", "").strip()

        options = GenerateOptions(
            task=task,
            words=words,
            language=language,
            notes=notes,
            output=output,
        )

        # Persist minimal session state (keeps UI behavior)
        session["options"] = {
            "task": options.normalized_task(),
            "words": options.clamped_words(),
            "language": options.normalized_language(),
            "output": options.normalized_output(),
            "notes": options.notes,
        }

        # Build corpus
        corpus, _metas = FileAnalyzer.extract_corpus(app, max_chars=120_000)

        # Generate
        service = SummarizerService()
        result_text = service.generate(corpus, options)

        # Store only, DO NOT download here
        session["last_result_text"] = result_text

        flash("Generated. Review the result on the right, then use 'Save & Download'.", "success")
        return redirect(url_for("index"))

    @app.post("/export")
    def export():
        text = request.form.get("result_text", "").strip()
        if not text:
            flash("Nothing to export.", "warning")
            return redirect(url_for("index"))

        opts = session.get("options", {}) or {}
        output = (opts.get("output") or "txt").lower()

        options = GenerateOptions(
            task=opts.get("task", "summary"),
            words=int(opts.get("words", 800)),
            language=opts.get("language", "English"),
            notes=opts.get("notes", ""),
            output=output
        )
        base_name = build_base_filename(options)

        if output == "pdf":
            path = os.path.join(current_app.config["OUTPUT_FOLDER"], f"{base_name}.pdf")
            return _export_pdf(text, path, f"{base_name}.pdf")
        elif output == "docx":
            path = os.path.join(current_app.config["OUTPUT_FOLDER"], f"{base_name}.docx")
            return _export_docx(text, path, f"{base_name}.docx")
        else:
            return _export_txt(text, f"{base_name}.txt")

    @app.get("/health")
    def health():
        return jsonify({"ok": True, "time": datetime.utcnow().isoformat() + "Z"})

    # -----------------------------------------------------------------------------
    # Export helpers
    # -----------------------------------------------------------------------------

    def _export_txt(text: str, download_name: str):
        bio = io.BytesIO()
        bio.write(text.encode("utf-8"))
        bio.seek(0)
        session["last_result_path"] = None
        return send_file(
            bio,
            mimetype="text/plain; charset=utf-8",
            as_attachment=True,
            download_name=download_name,
        )

    def _export_docx(text: str, path: str, download_name: str):
        if Document is None:
            return _export_txt(text, download_name.replace(".docx", ".txt"))

        doc = Document()
        for block in text.split("\n\n"):
            p = doc.add_paragraph()
            for line in block.splitlines():
                p.add_run(line)
                p.add_run("\n")
        doc.save(path)
        session["last_result_path"] = path
        return send_file(
            path,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=download_name,
        )

    def _export_pdf(text: str, path: str, download_name: str):
        if canvas is None:
            return _export_txt(text, download_name.replace(".pdf", ".txt"))

        c = canvas.Canvas(path)
        _register_pdf_font_if_available(c)
        try:
            c.setFont("DejaVuSans", 11)
        except Exception:
            pass
        _pdf_write_multiline(c, text, x=50, y_start=800, line_height=16)
        c.showPage()
        c.save()
        session["last_result_path"] = path
        return send_file(
            path,
            mimetype="application/pdf",
            as_attachment=True,
            download_name=download_name,
        )

    return app


# -----------------------------------------------------------------------------
# Entrypoint
# -----------------------------------------------------------------------------

if __name__ == "__main__":
    # Enable host/port override via env
    host = os.getenv("FLASK_RUN_HOST", "0.0.0.0")
    port = int(os.getenv("FLASK_RUN_PORT", "5000"))
    debug = os.getenv("FLASK_DEBUG", "false").lower() == "true"
    create_app().run(host=host, port=port, debug=debug)
